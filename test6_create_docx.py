from docx import Document
doc = Document()

title = doc.add_heading('Title', 0)  # 添加 word 文件的 title 标题

para = doc.add_paragraph('1. Personal Info')

doc.save('docs/test1.docx')

